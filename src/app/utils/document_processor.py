import io
import os
from typing import Dict, Any, Optional, Tuple
import logging
from pathlib import Path

# Document processing imports
import pypdf
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing utility for extracting text from various formats"""

    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt"]
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        file_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process document and extract text content

        Args:
            file_content: Raw file bytes
            filename: Original filename
            file_type: Optional file type hint

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Validate file size
            if len(file_content) > self.max_file_size:
                raise ValueError(f"File size exceeds limit of {self.max_file_size / (1024*1024):.1f}MB")

            # Determine file type
            file_extension = Path(filename).suffix.lower()
            if file_type:
                detected_type = file_type.lower()
            else:
                detected_type = self._detect_file_type(file_content, file_extension)

            logger.info(f"Processing document: {filename} (type: {detected_type})")

            # Process based on type
            if detected_type == "pdf":
                result = await self._process_pdf(file_content, filename)
            elif detected_type in ["docx", "doc"]:
                result = await self._process_docx(file_content, filename)
            elif detected_type == "txt":
                result = await self._process_text(file_content, filename)
            else:
                raise ValueError(f"Unsupported file type: {detected_type}")

            # Add common metadata
            result["metadata"].update({
                "original_filename": filename,
                "file_size": len(file_content),
                "detected_type": detected_type,
                "processing_status": "success"
            })

            logger.info(f"Successfully processed document: {filename}")
            return result

        except Exception as e:
            logger.error(f"Document processing failed for {filename}: {e}")
            return {
                "content": "",
                "metadata": {
                    "original_filename": filename,
                    "file_size": len(file_content),
                    "processing_status": "error",
                    "error": str(e)
                }
            }

    def _detect_file_type(self, file_content: bytes, file_extension: str) -> str:
        """
        Detect file type from content and extension

        Args:
            file_content: Raw file bytes
            file_extension: File extension from filename

        Returns:
            Detected file type
        """
        # Check magic bytes for common formats
        if file_content.startswith(b'%PDF'):
            return "pdf"
        elif file_content.startswith(b'PK'):  # ZIP-based formats like DOCX
            return "docx"
        elif file_extension in ['.pdf']:
            return "pdf"
        elif file_extension in ['.docx', '.doc']:
            return "docx"
        elif file_extension in ['.txt']:
            return "txt"
        else:
            # Try to decode as text
            try:
                file_content.decode('utf-8')
                return "txt"
            except UnicodeDecodeError:
                raise ValueError(f"Cannot determine file type for extension: {file_extension}")

    async def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process PDF document

        Args:
            file_content: PDF file bytes
            filename: Original filename

        Returns:
            Extracted content and metadata
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)

            # Extract text from all pages
            text_content = []
            page_count = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue

            full_text = "\n\n".join(text_content)

            # Extract metadata
            pdf_metadata = {}
            try:
                if pdf_reader.metadata:
                    pdf_metadata = {
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                        "creator": pdf_reader.metadata.get("/Creator", ""),
                        "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")),
                    }
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata: {e}")

            return {
                "content": full_text,
                "metadata": {
                    "page_count": page_count,
                    "character_count": len(full_text),
                    "word_count": len(full_text.split()),
                    "pdf_metadata": pdf_metadata
                }
            }

        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise

    async def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process DOCX document

        Args:
            file_content: DOCX file bytes
            filename: Original filename

        Returns:
            Extracted content and metadata
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract text from paragraphs and tables
            content_parts = []

            for element in doc.element.body:
                if isinstance(element, CT_P):  # Paragraph
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                elif isinstance(element, CT_Tbl):  # Table
                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        content_parts.append(f"[TABLE]\n{table_text}\n[/TABLE]")

            full_text = "\n\n".join(content_parts)

            # Extract document properties
            doc_properties = {}
            try:
                core_props = doc.core_properties
                doc_properties = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                }
            except Exception as e:
                logger.warning(f"Failed to extract DOCX properties: {e}")

            return {
                "content": full_text,
                "metadata": {
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables),
                    "character_count": len(full_text),
                    "word_count": len(full_text.split()),
                    "document_properties": doc_properties
                }
            }

        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise

    def _extract_table_text(self, table: Table) -> str:
        """
        Extract text from DOCX table

        Args:
            table: DOCX table object

        Returns:
            Formatted table text
        """
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_text.append(cell_text)
            table_text.append(" | ".join(row_text))
        return "\n".join(table_text)

    async def _process_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process plain text document

        Args:
            file_content: Text file bytes
            filename: Original filename

        Returns:
            Extracted content and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']
            text_content = None

            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                raise ValueError("Could not decode text file with any supported encoding")

            # Basic text processing
            lines = text_content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]

            return {
                "content": text_content,
                "metadata": {
                    "line_count": len(lines),
                    "non_empty_lines": len(non_empty_lines),
                    "character_count": len(text_content),
                    "word_count": len(text_content.split()),
                    "encoding_used": encoding
                }
            }

        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise

    def validate_document(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate document before processing

        Args:
            filename: Original filename
            file_size: File size in bytes

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.supported_formats:
            return False, f"Unsupported file format: {file_extension}"

        # Check file size
        if file_size > self.max_file_size:
            max_mb = self.max_file_size / (1024 * 1024)
            current_mb = file_size / (1024 * 1024)
            return False, f"File size {current_mb:.1f}MB exceeds limit of {max_mb:.1f}MB"

        # Check filename
        if not filename or len(filename) > 255:
            return False, "Invalid filename"

        return True, "Valid document"

    def get_processing_info(self) -> Dict[str, Any]:
        """
        Get information about document processing capabilities

        Returns:
            Processing capabilities information
        """
        return {
            "supported_formats": self.supported_formats,
            "max_file_size_mb": self.max_file_size / (1024 * 1024),
            "features": {
                "pdf": "Text extraction from all pages, metadata extraction",
                "docx": "Paragraph and table text extraction, document properties",
                "txt": "Multiple encoding support, basic text statistics"
            }
        }


# Global document processor instance
document_processor = DocumentProcessor()